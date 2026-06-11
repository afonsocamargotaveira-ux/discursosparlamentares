"""
Buscador de Discursos Parlamentares
------------------------------------
App em Streamlit que consulta a API de Dados Abertos da Câmara dos Deputados
(https://dadosabertos.camara.leg.br) e permite buscar discursos por assunto
(palavra-chave), exibindo o texto na íntegra.

Como rodar:
    pip install streamlit requests python-docx
    streamlit run app.py
"""

import streamlit as st
import requests
import time
from io import BytesIO
from datetime import date, timedelta
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docx.shared import Pt

API_BASE = "https://dadosabertos.camara.leg.br/api/v2"
HEADERS = {"Accept": "application/json"}

st.set_page_config(
    page_title="Discursos Parlamentares",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ----------------------------------------------------------------------
# Funções auxiliares (com cache para não sobrecarregar a API)
# ----------------------------------------------------------------------

@st.cache_data(ttl=86400, show_spinner=False)
def listar_legislaturas():
    params = {"itens": 100}
    legislaturas = []
    url = f"{API_BASE}/legislaturas"
    while url:
        resp = requests.get(url, params=params, headers=HEADERS, timeout=30)
        resp.raise_for_status()
        data = resp.json()
        legislaturas.extend(data["dados"])

        url = None
        params = None
        for link in data.get("links", []):
            if link["rel"] == "next":
                url = link["href"]

    legislaturas.sort(key=lambda x: x["id"], reverse=True)

    # Garantia: a 57ª legislatura (atual) às vezes não é retornada pela API
    if not any(leg["id"] == 57 for leg in legislaturas):
        legislaturas.insert(0, {"id": 57, "dataInicio": "2023-02-01", "dataFim": None})

    return legislaturas


@st.cache_data(ttl=3600, show_spinner=False)
def listar_deputados(legislatura=None, partido=None, uf=None, nome=None):
    params = {"itens": 100, "ordem": "ASC", "ordenarPor": "nome"}
    if legislatura:
        params["idLegislatura"] = legislatura
    if partido:
        params["siglaPartido"] = partido
    if uf:
        params["siglaUf"] = uf
    if nome:
        params["nome"] = nome

    deputados = []
    url = f"{API_BASE}/deputados"
    while url:
        resp = requests.get(url, params=params, headers=HEADERS, timeout=30)
        resp.raise_for_status()
        data = resp.json()
        deputados.extend(data["dados"])

        # paginação via links
        url = None
        params = None
        for link in data.get("links", []):
            if link["rel"] == "next":
                url = link["href"]
    # Deduplicar por id (a paginação ordenada por nome pode repetir registros
    # quando há empates de nome entre páginas)
    vistos = set()
    unicos = []
    for dep in deputados:
        if dep["id"] not in vistos:
            vistos.add(dep["id"])
            unicos.append(dep)
    return unicos


@st.cache_data(ttl=1800, show_spinner=False)
def listar_discursos(id_deputado, data_inicio, data_fim):
    params = {
        "dataInicio": data_inicio.isoformat(),
        "dataFim": data_fim.isoformat(),
        "ordenarPor": "dataHoraInicio",
        "ordem": "DESC",
        "itens": 100,
    }
    url = f"{API_BASE}/deputados/{id_deputado}/discursos"
    discursos = []
    while url:
        resp = requests.get(url, params=params, headers=HEADERS, timeout=15)
        if resp.status_code != 200:
            break
        data = resp.json()
        discursos.extend(data["dados"])

        url = None
        params = None
        for link in data.get("links", []):
            if link["rel"] == "next":
                url = link["href"]
    return discursos


@st.cache_data(ttl=1800, show_spinner=False)
def obter_texto_integral(url_texto):
    """Busca o HTML do discurso na íntegra e extrai o texto."""
    if not url_texto:
        return None
    try:
        resp = requests.get(url_texto, headers=HEADERS, timeout=10)
        resp.raise_for_status()
        # Conteúdo geralmente é HTML simples; extração leve sem bs4
        import re
        texto = resp.text
        texto = re.sub(r"<br\s*/?>", "\n", texto, flags=re.IGNORECASE)
        texto = re.sub(r"<[^>]+>", " ", texto)
        texto = re.sub(r"\n\s*\n+", "\n\n", texto)
        texto = re.sub(r"[ \t]+", " ", texto)
        return texto.strip()
    except requests.RequestException:
        return None


def classificar_local_discurso(discurso):
    """Classifica o discurso como 'Plenário' ou 'Comissão'.

    Discursos com faseEvento preenchida (ex: Pequeno Expediente, Grande
    Expediente, Ordem do Dia, Comunicações Parlamentares) ocorrem em
    sessões do Plenário. Os demais (reuniões de comissões, audiências
    públicas, etc.) são classificados como Comissão.
    """
    fase = discurso.get("faseEvento")
    if isinstance(fase, dict) and fase.get("titulo"):
        return "Plenário"
    return "Comissão"


def discurso_corresponde(discurso, termo):
    termo_lower = termo.lower()
    campos = [
        discurso.get("sumario", "") or "",
        discurso.get("keywords", "") or "",
        discurso.get("transcricao", "") or "",
        (discurso.get("faseEvento", {}).get("titulo") if isinstance(discurso.get("faseEvento"), dict) else "") or "",
    ]
    return any(termo_lower in c.lower() for c in campos)


def gerar_docx(resultados, termo_busca, textos_integrais):
    """Gera um arquivo .docx em memória com os discursos encontrados na íntegra."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    titulo = doc.add_heading("Discursos Parlamentares", level=0)
    doc.add_paragraph(f"Termo de busca: {termo_busca}")
    doc.add_paragraph(f"Total de discursos encontrados: {len(resultados)}")
    doc.add_paragraph("Fonte: API de Dados Abertos da Câmara dos Deputados (dadosabertos.camara.leg.br)")
    doc.add_page_break()

    for idx, (dep, d) in enumerate(resultados):
        data_hora = d.get("dataHoraInicio", "")
        data_fmt = data_hora[:10] if data_hora else "data não informada"

        doc.add_heading(
            f"{dep['nome']} ({dep.get('siglaPartido', '')}-{dep.get('siglaUf', '')}) — {data_fmt}",
            level=1,
        )

        if d.get("sumario"):
            p = doc.add_paragraph()
            p.add_run("Sumário: ").bold = True
            p.add_run(d.get("sumario", ""))

        if d.get("keywords"):
            p = doc.add_paragraph()
            p.add_run("Palavras-chave: ").bold = True
            p.add_run(d.get("keywords", ""))

        if d.get("tipoDiscurso"):
            p = doc.add_paragraph()
            p.add_run("Tipo: ").bold = True
            p.add_run(d.get("tipoDiscurso", ""))

        url_texto = d.get("urlTexto")
        if url_texto:
            p = doc.add_paragraph()
            p.add_run("Fonte original: ").bold = True
            p.add_run(url_texto)

        doc.add_paragraph("Texto na íntegra:").runs[0].bold = True

        texto_integral = textos_integrais.get(idx) or d.get("transcricao") or "Texto integral não disponível."
        for paragrafo in texto_integral.split("\n"):
            if paragrafo.strip():
                doc.add_paragraph(paragrafo.strip())

        if idx < len(resultados) - 1:
            doc.add_page_break()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ----------------------------------------------------------------------
# Interface
# ----------------------------------------------------------------------

# ----------------------------------------------------------------------
# Identidade visual
# ----------------------------------------------------------------------

st.markdown(
    """
    <style>
        /* Fundo geral */
        .stApp {
            background-color: #FFFFFF;
        }

        /* Sidebar */
        section[data-testid="stSidebar"] {
            background-color: #F7F8FA;
            border-right: 1px solid #E6E8EB;
        }
        section[data-testid="stSidebar"] h1,
        section[data-testid="stSidebar"] h2,
        section[data-testid="stSidebar"] h3 {
            color: #1B2A4A;
        }

        /* Tipografia geral */
        html, body, [class*="css"] {
            font-family: "Source Sans Pro", "Segoe UI", Helvetica, Arial, sans-serif;
        }

        /* Header / logo */
        .dp-header {
            display: flex;
            align-items: center;
            gap: 18px;
            padding: 18px 0 6px 0;
            border-bottom: 1px solid #E6E8EB;
            margin-bottom: 24px;
        }
        .dp-header h1 {
            font-size: 1.9rem;
            font-weight: 700;
            color: #1B2A4A;
            letter-spacing: 0.5px;
            margin: 0;
            line-height: 1.1;
        }
        .dp-header p {
            font-size: 0.95rem;
            color: #6B7280;
            margin: 2px 0 0 0;
        }

        /* Botões primários */
        .stButton > button[kind="primary"],
        .stDownloadButton > button[kind="primary"] {
            background-color: #1B2A4A;
            border-color: #1B2A4A;
            color: #FFFFFF;
            font-weight: 600;
            border-radius: 6px;
        }
        .stButton > button[kind="primary"]:hover,
        .stDownloadButton > button[kind="primary"]:hover {
            background-color: #2E4570;
            border-color: #2E4570;
        }

        /* Expanders (cards de discurso) */
        details[data-testid="stExpander"] {
            border: 1px solid #E6E8EB;
            border-radius: 8px;
            background-color: #FAFBFC;
            margin-bottom: 10px;
        }
        details[data-testid="stExpander"] summary {
            font-weight: 600;
            color: #1B2A4A;
        }

        /* Inputs */
        .stTextInput > div > div > input,
        .stTextArea textarea,
        .stSelectbox > div > div {
            border-radius: 6px;
        }

        footer {visibility: hidden;}
    </style>

    <div class="dp-header">
        <svg width="52" height="52" viewBox="0 0 52 52" xmlns="http://www.w3.org/2000/svg">
            <rect x="2" y="2" width="48" height="48" rx="8" fill="#1B2A4A"/>
            <path d="M26 11 L41 19 H11 Z" fill="#FFFFFF"/>
            <rect x="14" y="22" width="4" height="14" fill="#FFFFFF"/>
            <rect x="22" y="22" width="4" height="14" fill="#FFFFFF"/>
            <rect x="30" y="22" width="4" height="14" fill="#FFFFFF"/>
            <rect x="11" y="38" width="30" height="3" fill="#FFFFFF"/>
            <rect x="9" y="42" width="34" height="3" fill="#C9A227"/>
        </svg>
        <div>
            <h1>Discursos Parlamentares</h1>
            <p>Pesquisa de discursos do Plenário da Câmara dos Deputados — Dados Abertos</p>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

with st.sidebar:
    st.header("Filtros de busca")

    termo_busca = st.text_input(
        "Termo / assunto do discurso",
        placeholder="ex: educação, meio ambiente, reforma tributária...",
    )

    st.markdown("---")
    st.subheader("Local do discurso")
    locais_selecionados = st.multiselect(
        "Onde o discurso foi proferido:",
        options=["Plenário", "Comissão"],
        default=["Plenário", "Comissão"],
        help="Discursos vinculados a uma fase de sessão (ex: Pequeno/Grande "
        "Expediente, Ordem do Dia) são classificados como Plenário; os "
        "demais (reuniões de comissões, audiências públicas etc.) como "
        "Comissão.",
    )

    st.markdown("---")
    st.subheader("Período")

    modo_periodo = st.radio(
        "Definir período por:",
        options=["Data", "Legislatura"],
        horizontal=True,
    )

    if modo_periodo == "Data":
        col1, col2 = st.columns(2)
        with col1:
            data_inicio = st.date_input(
                "Data inicial", value=date.today() - timedelta(days=30)
            )
        with col2:
            data_fim = st.date_input("Data final", value=date.today())
        legislatura_escolhida = None
    else:
        try:
            legislaturas = listar_legislaturas()
        except requests.RequestException as e:
            st.error(f"Erro ao consultar legislaturas: {e}")
            st.stop()

        opcoes_legislatura = {"Todas as legislaturas": "todas"}
        for leg in legislaturas:
            inicio = (leg.get("dataInicio") or "")[:4]
            fim = (leg.get("dataFim") or "")[:4] or "atual"
            opcoes_legislatura[f"{leg['id']}ª Legislatura ({inicio}–{fim})"] = leg

        escolha = st.selectbox("Legislatura", options=list(opcoes_legislatura.keys()))
        legislatura_escolhida = opcoes_legislatura[escolha]
        data_inicio = None
        data_fim = None

    st.markdown("---")
    st.subheader("Filtrar deputados (opcional)")
    nome_dep = st.text_input("Nome do deputado", placeholder="ex: João Silva")
    partido = st.text_input("Sigla do partido", placeholder="ex: PT, PL, MDB...")
    uf = st.text_input("UF", placeholder="ex: SP, RJ, MG...")

    st.caption(
        "⚠️ A busca varre os discursos de **todos** os deputados que se "
        "enquadrem nos filtros acima, dentro do período/legislatura "
        "escolhido. Sem filtros adicionais, e principalmente na opção "
        "**'Todas as legislaturas'**, isso pode envolver milhares de "
        "deputados/períodos e levar bastante tempo."
    )

    if st.button("🔍 Buscar discursos", type="primary", use_container_width=True):
        st.session_state["buscar_disparado"] = True
        st.session_state.pop("resultados", None)


if not st.session_state.get("buscar_disparado"):
    st.info(
        "Defina o termo de busca, o período e (opcionalmente) filtros de "
        "deputado/partido/UF na barra lateral e clique em **Buscar discursos**."
    )
    st.stop()

if not termo_busca.strip():
    st.warning("Informe um termo de busca para continuar.")
    st.stop()

if not locais_selecionados:
    st.warning("Selecione ao menos um local de discurso (Plenário e/ou Comissão).")
    st.stop()

if modo_periodo == "Data" and data_inicio > data_fim:
    st.error("A data inicial não pode ser posterior à data final.")
    st.stop()

# ----------------------------------------------------------------------
# Execução da busca (apenas se ainda não houver resultados em cache)
# ----------------------------------------------------------------------

if "resultados" not in st.session_state:

    # tarefas: lista de tuplas (deputado, data_inicio, data_fim)
    tarefas = []

    with st.spinner("Carregando lista de deputados..."):
        try:
            if modo_periodo == "Data":
                deputados = listar_deputados(
                    partido=partido.strip() or None,
                    uf=uf.strip().upper() or None,
                    nome=nome_dep.strip() or None,
                )
                for dep in deputados:
                    tarefas.append((dep, data_inicio, data_fim))

            elif legislatura_escolhida == "todas":
                legislaturas = listar_legislaturas()
                vistos = set()
                for leg in legislaturas:
                    leg_inicio = leg.get("dataInicio")
                    if not leg_inicio:
                        continue
                    d_ini = date.fromisoformat(leg_inicio[:10])
                    leg_fim = leg.get("dataFim")
                    d_fim = date.fromisoformat(leg_fim[:10]) if leg_fim else date.today()
                    deps_leg = listar_deputados(
                        legislatura=leg["id"],
                        partido=partido.strip() or None,
                        uf=uf.strip().upper() or None,
                        nome=nome_dep.strip() or None,
                    )
                    for dep in deps_leg:
                        chave = (dep["id"], leg["id"])
                        if chave in vistos:
                            continue
                        vistos.add(chave)
                        tarefas.append((dep, d_ini, d_fim))

            else:
                leg = legislatura_escolhida
                d_ini = date.fromisoformat(leg["dataInicio"][:10])
                d_fim = date.fromisoformat(leg["dataFim"][:10]) if leg.get("dataFim") else date.today()
                deputados = listar_deputados(
                    legislatura=leg["id"],
                    partido=partido.strip() or None,
                    uf=uf.strip().upper() or None,
                    nome=nome_dep.strip() or None,
                )
                for dep in deputados:
                    tarefas.append((dep, d_ini, d_fim))

        except requests.RequestException as e:
            st.error(f"Erro ao consultar a API da Câmara: {e}")
            st.stop()

    if not tarefas:
        st.warning("Nenhum deputado encontrado com os filtros informados.")
        st.stop()

    st.write(f"Varrendo discursos de **{len(tarefas)}** combinação(ões) deputado/período...")

    progresso = st.progress(0.0)
    status = st.empty()

    resultados = []
    concluidas = 0

    def _buscar(tarefa):
        dep, d_ini, d_fim = tarefa
        try:
            discursos = listar_discursos(dep["id"], d_ini, d_fim)
        except Exception:
            discursos = []
        encontrados = [
            (dep, d)
            for d in discursos
            if discurso_corresponde(d, termo_busca)
            and classificar_local_discurso(d) in locais_selecionados
        ]
        return dep, encontrados

    with ThreadPoolExecutor(max_workers=8) as executor:
        futures = {executor.submit(_buscar, t): t for t in tarefas}
        for future in as_completed(futures):
            dep, encontrados = future.result()
            resultados.extend(encontrados)
            concluidas += 1
            status.text(
                f"Consultando discursos... ({concluidas}/{len(tarefas)}) "
                f"— última: {dep['nome']} — encontrados até agora: {len(resultados)}"
            )
            progresso.progress(concluidas / len(tarefas))

    status.empty()
    progresso.empty()

    resultados.sort(key=lambda r: r[1].get("dataHoraInicio", ""), reverse=True)
    st.session_state["resultados"] = resultados
    st.session_state["termo_busca"] = termo_busca

resultados = st.session_state["resultados"]
termo_busca = st.session_state.get("termo_busca", termo_busca)

# ----------------------------------------------------------------------
# Exibição dos resultados
# ----------------------------------------------------------------------

st.markdown("---")

if not resultados:
    st.warning(
        "Nenhum discurso encontrado com esse termo no período/deputados "
        "selecionados. Tente ampliar o período, aumentar o número de "
        "deputados varridos ou usar outro termo."
    )
    if st.button("🔄 Nova busca"):
        st.session_state.pop("buscar_disparado", None)
        st.session_state.pop("resultados", None)
        st.rerun()
    st.stop()

st.success(f"Foram encontrados **{len(resultados)}** discurso(s).")

st.markdown("---")

gerar_word = st.button(
    "📄 Gerar arquivo Word com os discursos na íntegra",
    type="primary",
    use_container_width=True,
)

if gerar_word:
    progresso_w = st.progress(0.0)
    status_w = st.empty()

    with ThreadPoolExecutor(max_workers=40) as executor:
        future_to_idx = {
            executor.submit(obter_texto_integral, d.get("urlTexto")): idx
            for idx, (dep, d) in enumerate(resultados)
        }
        textos_integrais = {}
        concluidos_w = 0
        for future in as_completed(future_to_idx):
            idx = future_to_idx[future]
            try:
                texto = future.result()
            except Exception:
                texto = None
            if not texto:
                texto = resultados[idx][1].get("transcricao") or "Texto integral não disponível para este discurso."
            textos_integrais[idx] = texto
            concluidos_w += 1
            if concluidos_w % 25 == 0 or concluidos_w == len(resultados):
                status_w.text(f"Carregando textos... ({concluidos_w}/{len(resultados)})")
                progresso_w.progress(concluidos_w / len(resultados))

    status_w.text("Montando o arquivo Word...")
    docx_buffer = gerar_docx(resultados, termo_busca, textos_integrais)
    st.session_state["docx_buffer"] = docx_buffer.getvalue()

    status_w.empty()
    progresso_w.empty()

if "docx_buffer" in st.session_state:
    st.download_button(
        label="✅ Baixar discursos em Word (.docx)",
        data=st.session_state["docx_buffer"],
        file_name=f"discursos_{termo_busca.strip().replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True,
    )

st.markdown("---")
if st.button("🔄 Nova busca"):
    for chave in ["buscar_disparado", "resultados", "termo_busca", "docx_buffer"]:
        st.session_state.pop(chave, None)
    st.rerun()
